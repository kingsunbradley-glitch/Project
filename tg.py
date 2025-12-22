import docx
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_chinese_academic_paper():
    doc = docx.Document()

    # --- 1. 页面设置 (A4 标准: 上下2.54cm, 左右3.17cm) ---
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # --- 2. 样式辅助函数：处理中西文混排字体 ---
    def set_run_font(run, font_size_pt, is_bold=False, is_heading=False):
        """
        设置中西文混排字体。
        西文: Times New Roman
        中文: 标题用黑体，正文用宋体
        """
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体' if is_heading else '宋体')
        run.font.size = Pt(font_size_pt)
        run.bold = is_bold

    # --- 3. 段落辅助函数 ---
    def add_paper_title(text):
        """添加论文主标题 (小二, 加粗, 居中, 黑体)"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(18)
        run = p.add_run(text)
        set_run_font(run, 18, is_bold=True, is_heading=True) # 小二 = 18pt

    def add_heading_1(text):
        """添加一级标题 (四号, 加粗, 左对齐, 黑体)"""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(text)
        set_run_font(run, 14, is_bold=True, is_heading=True) # 四号 = 14pt

    def add_heading_2(text):
        """添加二级标题 (小四, 加粗, 左对齐, 黑体)"""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(text)
        set_run_font(run, 12, is_bold=True, is_heading=True) # 小四 = 12pt

    def add_body_text(text):
        """添加正文 (小四, 宋体, 首行缩进2字符, 1.5倍行距)"""
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(24) # 缩进约2个字符 (12pt * 2)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY # 两端对齐
        run = p.add_run(text)
        set_run_font(run, 12, is_bold=False, is_heading=False) # 小四 = 12pt

    # ================= 写作内容开始 =================

    # 标题
    add_paper_title('273Ds 衰变性质研究')

    # 第一部分
    add_heading_1('1. 选题的背景及意义')
    
    add_heading_2('1.1 超重核结构与壳效应问题')
    add_body_text('超重元素区的存在与稳定性高度依赖于微观壳效应。理论预言在球形“稳定岛”（Z ~ 114–126, N ~ 184）之外，还存在由变形驱动的次级壳隙（Deformed Shell Gaps），尤其是 N=162 和 Z=108 区域。这些变形壳隙对原子核的基态性质（如质量、形变）及衰变性质（半衰期、Qα）具有决定性影响。')
    add_body_text('尽管目前的实验与理论模型普遍支持 N=162 变形壳的存在，但其强度、局限性以及模型依赖性仍存在较大争议。尤其是当质子数远离 Z=108 (Hs) 时，壳效应的衰减趋势尚需实验数据的精确约束。')

    add_heading_2('1.2 273Ds 及其衰变链的特殊地位')
    add_body_text('273Ds（Z=110, N=163）位于 N=162 壳隙的边缘，通过 238U + 40Ar 热熔合反应制备。其 α 衰变链 273Ds → 269Hs → 265Sg → 261Rf 能够直接穿越 N=162 壳隙区域，是探测该区域单粒子能级结构和同核异能态（Isomer）的理想探针。')

    add_heading_2('1.3 选题意义')
    add_body_text('本课题旨在通过高统计量的实验数据，厘清 273Ds 及其子核 269Hs、265Sg 的能级结构。既往实验（如 DGFRS, SHIP）在 273Ds 的衰变性质上存在显著分歧（如 GSI 的短寿命数据与 JINR 报道的长寿命 Isomer 数据之间的矛盾）。通过系统分析，本研究有望解决 269Hs 和 265Sg 中长期存在的 Isomer 疑题，为改进宏观-微观模型及相对论平均场理论在超重核区的预测能力提供关键实验依据。')

    # 第二部分
    add_heading_1('2. 国内外本学科领域的发展现状与趋势')
    
    add_heading_2('2.1 超重元素的合成与衰变研究现状')
    add_body_text('国际上，GSI（德国）、RIKEN（日本）、JINR（俄罗斯）等实验室已成功合成了一系列 Ds 同位素。对于 273Ds，早期 GSI 的 SHIP 实验和 RIKEN 的 GARIS 实验观测到了半衰期约为 0.17 ms 的短寿命组分。然而，JINR 的 Oganessian 等人在随后的 238U + 40Ar 实验中报道了半衰期长达 41 ms 的衰变链，并将其归因于 273Ds 的同核异能态。')
    add_body_text('目前的争议焦点在于：JINR 报道的长寿命组分，其子核 265Sg 的 α 衰变能量显著低于文献值，且 273Ds 的衰变时间比理论值高出近两个数量级。这种实验数据的“模糊性”亟待新的独立数据分析来澄清。')

    add_heading_2('2.2 Isomer 与单粒子结构研究进展')
    add_body_text('在 N ≈ 162 区域，由于高 j 轨道（如中子 11/2-[725] 等）的出现，奇 A 核中普遍存在同核异能态（Isomer）。文献已确认 265Sg 和 261Rf 存在两组 α 能量（a/b 态），但在能级排序（何者为基态）上尚无定论。对于 269Hs，虽有迹象表明存在 Isomer，但缺乏完整链条的直接观测证据。')

    # 第三部分
    add_heading_1('3. 课题主要研究内容、预期目标')
    
    add_heading_2('3.1 主要研究内容')
    add_body_text('1. 数据挖掘与整理：系统分析本实验（238U + 40Ar）获取的实验数据，重点分析两条特征迥异的衰变路径：')
    # 列表项通常悬挂缩进，这里简单处理为段落
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run('   - 路径 A（高 Qα、短寿命）：273Ds (~0.1 ms) → 269Hs → 265Sg\n   - 路径 B（低 Qα、长寿命）：273Ds (~8 ms) → 269Hs → 265Sg')
    set_run_font(run, 12)

    add_body_text('2. 269Hs Isomer 定量分析：结合本实验观测到的完整链条，对比文献中缺失 α 粒子的不完整链条，定量给出 269Hs 两个态的能量差、半衰期及分支比。')
    add_body_text('3. 265Sg/261Rf 双态结构再评价：利用 α 阻滞因子（Hindrance Factor）分析，判断各衰变级间的跃迁性质（是否发生自旋翻转），从而约束下游核素的基态/Isomer 指定。')

    add_heading_2('3.2 预期目标')
    add_body_text('1. 给出 273Ds 两条衰变分支的精确 Qα 值和半衰期 T1/2。')
    add_body_text('2. 确认 269Hs 的 Isomer 存在，并尝试建立包含 Ds-Hs-Sg-Rf 的能级系统图。')
    add_body_text('3. 评价 N=162 壳隙在 Z=110 处的强度，讨论现有理论模型（如宏观-微观模型）的适用性。')

    # 第四部分
    add_heading_1('4. 拟采用的研究方法、技术路线及可行性分析')
    
    add_heading_2('4.1 研究方法')
    add_body_text('1. 关联分析法：利用时间-位置关联技术（Time-Position Correlation），在 DSSD（双面硅条探测器）数据中构建 ER-α-α-SF 衰变链，严格排除随机本底。')
    add_body_text('2. 统计推断：鉴于超重核事件稀少，采用最大似然法（MLH）和贝叶斯统计方法提取半衰期及其不确定度，确保结果的统计学严谨性。')
    add_body_text('3. 理论辅助：计算 α 衰变阻滞因子（HF），结合 Nilsson 单粒子能级图，对观测到的跃迁进行组态指认（如 1/2+[620] 等轨道）。')
    
    add_heading_2('4.2 可行性分析')
    add_body_text('本课题依托已完成的物理实验，原始数据积累充足，无需重新申请束流时间。团队已开发成熟的 ROOT 分析框架，且已有文献数据（GSI/JINR）提供了丰富的对比基准，课题在技术和数据层面均具备高度可行性。')

    # 第五部分
    add_heading_1('5. 已有研究基础与所需的研究条件')
    add_body_text('1. 已获得数百 GB 的 238U + 40Ar 实验原始数据，初步扫描已发现多条完整的 273Ds 候选链。')
    add_body_text('2. 已完成 DSSD 探测器的精细能量刻度与分辨率校准，能量分辨率满足分辨 Isomer 的需求。')
    add_body_text('3. 具备高性能数据处理工作站及 ROOT、C++、Python 分析环境。')

    # 第六部分
    add_heading_1('6. 研究工作计划与进度安排')
    
    # 创建表格
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.autofit = True
    
    # 设置表格标题行
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '时间阶段'
    hdr_cells[1].text = '主要工作内容'
    # 加粗表头
    for cell in hdr_cells:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run_font(run, 12, is_bold=True, is_heading=True)

    # 填充数据
    plan_data = [
        ('第 1-3 个月', '数据复核与能量刻度优化；完成 273Ds 候选事件的初筛。'),
        ('第 4-6 个月', '构建精细衰变链，利用贝叶斯方法提取半衰期与分支比。'),
        ('第 7-9 个月', '进行核结构分析，计算阻滞因子，构建能级纲图。'),
        ('第 10-12 个月', '撰写研究报告与学术论文，准备结题或答辩。')
    ]

    for time_period, content in plan_data:
        row_cells = table.add_row().cells
        row_cells[0].text = time_period
        row_cells[1].text = content
        # 设置表格内容字体
        for cell in row_cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    set_run_font(run, 12, is_bold=False, is_heading=False)

    # 第七部分
    add_heading_1('7. 参考文献')
    
    refs = [
        "[1] Yu. Ts. Oganessian et al., Phys. Rev. C (2000/2004). (关于 273Ds 合成及 Isomer 争议的原始文献)",
        "[2] S. Hofmann et al. (GSI Collaboration), Eur. Phys. J. A. (SHIP 实验关于 273Ds 的早期结果)",
        "[3] K. Morita et al. (RIKEN Collaboration), J. Phys. Soc. Jpn. (GARIS 实验数据)",
        "[4] A. Türler et al., Nuclear structure and reaction studies near doubly magic 270Hs.",
        "[5] F.P. Hessberger, Eur. Phys. J. A (2017). (综述：超重核区的 Isomer 结构)",
        "[6] Audi, G., et al. (NUBASE2020), Chin. Phys. C (2021). (基准核数据)"
    ]

    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(0) # 参考文献通常悬挂缩进或不缩进
        p.paragraph_format.left_indent = Pt(0) 
        run = p.add_run(ref)
        set_run_font(run, 10.5) # 五号字 = 10.5pt

    # 保存文件
    filename = '273Ds_衰变性质研究_中文论文格式.docx'
    doc.save(filename)
    print(f"文档已生成: {filename}")

if __name__ == "__main__":
    create_chinese_academic_paper()